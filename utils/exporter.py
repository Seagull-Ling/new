# -*- coding: utf-8 -*-
"""
Excel/Word导出功能模块
实现座位表和明细表的导出功能
"""
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO
import copy

try:
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class Exporter:
    """导出器"""
    
    def __init__(self):
        pass
    
    def validate_data(self, people: List[Dict], seats: List[Dict], 
                       row_configs: List[Dict]) -> Dict[str, Any]:
        """
        导出前的数据校验
        
        校验内容：
        - 是否有人未安排座位
        - 是否存在空白姓名
        - 是否人数超过座位数
        - 是否存在重复姓名（提醒即可）
        
        Args:
            people: 人员列表
            seats: 座位列表
            row_configs: 每排配置
            
        Returns:
            Dict: 校验结果，包含errors和warnings
        """
        errors = []
        warnings = []
        
        # 统计已分配和未分配座位的人员
        assigned_count = sum(1 for p in people if p.get('current_seat_id'))
        unassigned_count = len(people) - assigned_count
        
        # 检查是否有人未安排座位
        if unassigned_count > 0:
            warnings.append(f"有 {unassigned_count} 人未安排座位")
        
        # 检查是否存在空白姓名
        blank_names = [p for p in people if not p.get('name') or str(p.get('name', '')).strip() == '']
        if blank_names:
            warnings.append(f"存在 {len(blank_names)} 个空白姓名")
        
        # 检查是否人数超过座位数
        total_seats = len(seats)
        if len(people) > total_seats and total_seats > 0:
            errors.append(f"人数({len(people)})超过座位数({total_seats})")
        
        # 检查是否存在重复姓名
        name_counts = {}
        for p in people:
            name = p.get('name', '')
            if name:
                name_counts[name] = name_counts.get(name, 0) + 1
        
        duplicate_names = [name for name, count in name_counts.items() if count > 1]
        if duplicate_names:
            warnings.append(f"存在重复姓名: {', '.join(duplicate_names[:5])}" + 
                          ("等" if len(duplicate_names) > 5 else ""))
        
        return {
            'valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings,
            'stats': {
                'total_people': len(people),
                'assigned_people': assigned_count,
                'unassigned_people': unassigned_count,
                'total_seats': total_seats,
                'duplicate_names': duplicate_names
            }
        }
    
    def export_excel(self, people: List[Dict], seats: List[Dict], 
                     row_configs: List[Dict]) -> Tuple[BytesIO, Dict[str, Any]]:
        """
        导出Excel文件
        
        包含3个工作表：
        1. 正式座位表：以排座视图输出，每排显示左右座位与中间分界
        2. 当前姓名顺序表：当前顺序号、姓名、当前座位
        3. 座位分配明细表：全局顺序号、排号、左右侧、座位编号、显示名称、姓名
        
        Args:
            people: 人员列表
            seats: 座位列表
            row_configs: 每排配置
            
        Returns:
            Tuple[BytesIO, Dict]: (Excel文件字节流, 校验结果)
        """
        if not HAS_OPENPYXL:
            raise ImportError("请安装openpyxl和pandas: pip install openpyxl pandas")
        
        # 先校验数据
        validation = self.validate_data(people, seats, row_configs)
        
        # 创建工作簿
        wb = Workbook()
        
        # 删除默认的Sheet
        default_sheet = wb.active
        wb.remove(default_sheet)
        
        # 工作表1：正式座位表
        ws1 = wb.create_sheet("正式座位表", 0)
        self._create_seating_chart_sheet(ws1, people, seats, row_configs)
        
        # 工作表2：当前姓名顺序表
        ws2 = wb.create_sheet("当前姓名顺序表", 1)
        self._create_name_order_sheet(ws2, people, seats)
        
        # 工作表3：座位分配明细表
        ws3 = wb.create_sheet("座位分配明细表", 2)
        self._create_detail_sheet(ws3, people, seats)
        
        # 保存到字节流
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        return output, validation
    
    def _create_seating_chart_sheet(self, ws, people: List[Dict], 
                                    seats: List[Dict], row_configs: List[Dict]):
        """
        创建正式座位表工作表
        """
        # 定义样式
        header_font = Font(bold=True, size=12)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=12, color="FFFFFF")
        center_align = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 标题
        ws['A1'] = '正式座位表'
        ws['A1'].font = Font(bold=True, size=16)
        ws.merge_cells('A1:Z1')
        ws['A1'].alignment = center_align
        
        row_idx = 3
        
        # 按排号顺序处理
        row_nos = sorted(set(s['row_no'] for s in seats))
        
        for row_no in row_nos:
            # 获取该排的物理顺序座位
            row_seats = [s for s in seats if s['row_no'] == row_no]
            
            # 分离左右侧
            left_seats = [s for s in row_seats if s['side'] == 'LEFT']
            right_seats = [s for s in row_seats if s['side'] == 'RIGHT']
            
            # 排序：左侧倒序（左4,左3,左2,左1），右侧正序（右1,右2,右3,右4）
            left_sorted = sorted(left_seats, key=lambda x: x['seat_no'], reverse=True)
            right_sorted = sorted(right_seats, key=lambda x: x['seat_no'])
            
            # 排标题
            title_cell = ws.cell(row=row_idx, column=1, value=f'第{row_no}排')
            title_cell.font = header_font
            title_cell.alignment = center_align
            
            # 计算需要的列数
            total_cols = len(left_sorted) + 1 + len(right_sorted)  # +1是中间分隔
            
            # 合并排标题单元格
            if total_cols > 1:
                ws.merge_cells(start_row=row_idx, start_column=1, 
                               end_row=row_idx, end_column=total_cols)
            
            row_idx += 1
            
            # 写入座位标签行
            col_idx = 1
            
            # 左侧座位
            for seat in left_sorted:
                cell = ws.cell(row=row_idx, column=col_idx, value=seat['display_label'])
                cell.font = header_font_white
                cell.fill = header_fill
                cell.alignment = center_align
                cell.border = thin_border
                col_idx += 1
            
            # 中间分隔
            center_cell = ws.cell(row=row_idx, column=col_idx, value='【中间】')
            center_cell.font = Font(bold=True, size=10, color="FF0000")
            center_cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
            center_cell.alignment = center_align
            center_cell.border = thin_border
            col_idx += 1
            
            # 右侧座位
            for seat in right_sorted:
                cell = ws.cell(row=row_idx, column=col_idx, value=seat['display_label'])
                cell.font = header_font_white
                cell.fill = header_fill
                cell.alignment = center_align
                cell.border = thin_border
                col_idx += 1
            
            row_idx += 1
            
            # 写入姓名行
            col_idx = 1
            
            # 左侧座位姓名
            for seat in left_sorted:
                person = self._get_person_by_seat(seat, people)
                name = person['name'] if person else ''
                cell = ws.cell(row=row_idx, column=col_idx, value=name)
                cell.alignment = center_align
                cell.border = thin_border
                col_idx += 1
            
            # 中间分隔
            center_cell2 = ws.cell(row=row_idx, column=col_idx, value='')
            center_cell2.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
            center_cell2.alignment = center_align
            center_cell2.border = thin_border
            col_idx += 1
            
            # 右侧座位姓名
            for seat in right_sorted:
                person = self._get_person_by_seat(seat, people)
                name = person['name'] if person else ''
                cell = ws.cell(row=row_idx, column=col_idx, value=name)
                cell.alignment = center_align
                cell.border = thin_border
                col_idx += 1
            
            row_idx += 2  # 排之间空一行
        
        # 调整列宽
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 20)
            ws.column_dimensions[column].width = adjusted_width
    
    def _create_name_order_sheet(self, ws, people: List[Dict], seats: List[Dict]):
        """
        创建当前姓名顺序表工作表
        """
        # 定义样式
        header_font = Font(bold=True, size=11)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=11, color="FFFFFF")
        center_align = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 标题
        ws['A1'] = '当前姓名顺序表'
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:C1')
        ws['A1'].alignment = center_align
        
        # 表头
        headers = ['当前顺序号', '姓名', '当前座位']
        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=3, column=col_idx, value=header)
            cell.font = header_font_white
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
        
        # 数据
        # 按order_index排序
        sorted_people = sorted(people, key=lambda x: x['order_index'])
        
        for row_idx, person in enumerate(sorted_people, start=4):
            # 顺序号
            cell1 = ws.cell(row=row_idx, column=1, value=person['order_index'] + 1)
            cell1.alignment = center_align
            cell1.border = thin_border
            
            # 姓名
            cell2 = ws.cell(row=row_idx, column=2, value=person['name'])
            cell2.alignment = center_align
            cell2.border = thin_border
            
            # 当前座位
            seat_id = person.get('current_seat_id')
            seat = self._get_seat_by_id(seat_id, seats)
            seat_label = seat['display_label'] if seat else '未安排'
            cell3 = ws.cell(row=row_idx, column=3, value=seat_label)
            cell3.alignment = center_align
            cell3.border = thin_border
        
        # 调整列宽
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
    
    def _create_detail_sheet(self, ws, people: List[Dict], seats: List[Dict]):
        """
        创建座位分配明细表工作表
        """
        # 定义样式
        header_font = Font(bold=True, size=11)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=11, color="FFFFFF")
        center_align = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 标题
        ws['A1'] = '座位分配明细表'
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:F1')
        ws['A1'].alignment = center_align
        
        # 表头
        headers = ['全局顺序号', '排号', '左右侧', '座位编号', '显示名称', '姓名']
        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=3, column=col_idx, value=header)
            cell.font = header_font_white
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
        
        # 按排座顺序排序
        sorted_seats = sorted(seats, key=lambda x: x['seat_order'])
        
        for row_idx, seat in enumerate(sorted_seats, start=4):
            # 全局顺序号
            cell1 = ws.cell(row=row_idx, column=1, value=seat['seat_order'] + 1)
            cell1.alignment = center_align
            cell1.border = thin_border
            
            # 排号
            cell2 = ws.cell(row=row_idx, column=2, value=f"第{seat['row_no']}排")
            cell2.alignment = center_align
            cell2.border = thin_border
            
            # 左右侧
            side_text = '右侧' if seat['side'] == 'RIGHT' else '左侧'
            cell3 = ws.cell(row=row_idx, column=3, value=side_text)
            cell3.alignment = center_align
            cell3.border = thin_border
            
            # 座位编号
            cell4 = ws.cell(row=row_idx, column=4, value=seat['seat_no'])
            cell4.alignment = center_align
            cell4.border = thin_border
            
            # 显示名称
            cell5 = ws.cell(row=row_idx, column=5, value=seat['display_label'])
            cell5.alignment = center_align
            cell5.border = thin_border
            
            # 姓名
            person = self._get_person_by_seat(seat, people)
            name = person['name'] if person else '（空）'
            cell6 = ws.cell(row=row_idx, column=6, value=name)
            cell6.alignment = center_align
            cell6.border = thin_border
        
        # 调整列宽
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 10
        ws.column_dimensions['C'].width = 8
        ws.column_dimensions['D'].width = 10
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 15
    
    def _get_person_by_seat(self, seat: Dict, people: List[Dict]) -> Optional[Dict]:
        """
        根据座位获取分配的人员
        """
        person_id = seat.get('assigned_person_id')
        if not person_id:
            return None
        for p in people:
            if p.get('person_id') == person_id:
                return p
        return None
    
    def _get_seat_by_id(self, seat_id: str, seats: List[Dict]) -> Optional[Dict]:
        """
        根据seat_id获取座位
        """
        if not seat_id:
            return None
        for s in seats:
            if s.get('seat_id') == seat_id:
                return s
        return None
    
    def export_word(self, people: List[Dict], seats: List[Dict],
                    row_configs: List[Dict]) -> Tuple[BytesIO, Dict[str, Any]]:
        """
        导出Word文件
        
        包含：
        1. 正式座位表
        2. 座位分配明细表
        
        Args:
            people: 人员列表
            seats: 座位列表
            row_configs: 每排配置
            
        Returns:
            Tuple[BytesIO, Dict]: (Word文件字节流, 校验结果)
        """
        if not HAS_DOCX:
            raise ImportError("请安装python-docx: pip install python-docx")
        
        # 先校验数据
        validation = self.validate_data(people, seats, row_configs)
        
        # 创建文档
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 标题
        title = doc.add_heading('会议座位安排表', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 第一部分：正式座位表
        doc.add_heading('一、正式座位表', level=1)
        
        # 按排号顺序处理
        row_nos = sorted(set(s['row_no'] for s in seats))
        
        for row_no in row_nos:
            # 排标题
            row_title = doc.add_paragraph()
            row_title.add_run(f'第{row_no}排').bold = True
            
            # 获取该排的物理顺序座位
            row_seats = [s for s in seats if s['row_no'] == row_no]
            
            # 分离左右侧
            left_seats = [s for s in row_seats if s['side'] == 'LEFT']
            right_seats = [s for s in row_seats if s['side'] == 'RIGHT']
            
            # 排序
            left_sorted = sorted(left_seats, key=lambda x: x['seat_no'], reverse=True)
            right_sorted = sorted(right_seats, key=lambda x: x['seat_no'])
            
            # 创建表格：座位标签行 + 姓名行
            total_cols = len(left_sorted) + 1 + len(right_sorted)
            table = doc.add_table(rows=2, cols=total_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # 第一行：座位标签
            col_idx = 0
            
            # 左侧座位标签
            for seat in left_sorted:
                cell = table.rows[0].cells[col_idx]
                cell.text = seat['display_label']
                # 设置样式
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                col_idx += 1
            
            # 中间分隔
            center_cell = table.rows[0].cells[col_idx]
            center_cell.text = '【中间】'
            for para in center_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
            col_idx += 1
            
            # 右侧座位标签
            for seat in right_sorted:
                cell = table.rows[0].cells[col_idx]
                cell.text = seat['display_label']
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                col_idx += 1
            
            # 第二行：姓名
            col_idx = 0
            
            # 左侧座位姓名
            for seat in left_sorted:
                person = self._get_person_by_seat(seat, people)
                name = person['name'] if person else ''
                cell = table.rows[1].cells[col_idx]
                cell.text = name
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                col_idx += 1
            
            # 中间分隔
            center_cell2 = table.rows[1].cells[col_idx]
            center_cell2.text = ''
            col_idx += 1
            
            # 右侧座位姓名
            for seat in right_sorted:
                person = self._get_person_by_seat(seat, people)
                name = person['name'] if person else ''
                cell = table.rows[1].cells[col_idx]
                cell.text = name
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                col_idx += 1
            
            doc.add_paragraph()  # 空行
        
        # 第二部分：座位分配明细表
        doc.add_heading('二、座位分配明细表', level=1)
        
        # 创建表格
        sorted_seats = sorted(seats, key=lambda x: x['seat_order'])
        
        # 表头：全局顺序号、排号、左右侧、座位编号、显示名称、姓名
        detail_table = doc.add_table(rows=len(sorted_seats) + 1, cols=6)
        detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        detail_table.style = 'Table Grid'
        
        # 表头行
        headers = ['全局顺序号', '排号', '左右侧', '座位编号', '显示名称', '姓名']
        for col_idx, header in enumerate(headers):
            cell = detail_table.rows[0].cells[col_idx]
            cell.text = header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
        
        # 数据行
        for row_idx, seat in enumerate(sorted_seats, start=1):
            row = detail_table.rows[row_idx]
            
            # 全局顺序号
            row.cells[0].text = str(seat['seat_order'] + 1)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 排号
            row.cells[1].text = f"第{seat['row_no']}排"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 左右侧
            side_text = '右侧' if seat['side'] == 'RIGHT' else '左侧'
            row.cells[2].text = side_text
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 座位编号
            row.cells[3].text = str(seat['seat_no'])
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 显示名称
            row.cells[4].text = seat['display_label']
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 姓名
            person = self._get_person_by_seat(seat, people)
            name = person['name'] if person else '（空）'
            row.cells[5].text = name
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存到字节流
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output, validation
